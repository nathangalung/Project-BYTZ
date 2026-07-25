import json
import logging
import os
from urllib.parse import urlparse
from collections.abc import AsyncIterator
from time import perf_counter
from typing import Literal

import httpx
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from app.middleware.auth import require_service_auth
from app.models.schemas import (
    BrdSectionScore,
    BrdTemplateScore,
    ChatRequest,
    ChatResponse,
    CertificationEntry,
    CvParseRequest,
    CvParseResponse,
    CvParsedData,
    EducationEntry,
    ExperienceEntry,
    GenerateBrdRequest,
    GenerateBrdResponse,
    GeneratePrdRequest,
    GeneratePrdResponse,
    MatchingRequest,
    ProjectEntry,
    MatchingResponse,
    ParseSpecData,
    ParseSpecRequest,
    ParseSpecResponse,
)
from app.services.llm import (
    LLMError,
    LlmUsage,
    generate_json,
    generate_structured,
    generate_text,
    stream_text,
)
from app.services.nats_client import publish_event
from app.services.usage import InteractionStatus, record_interaction, track

logger = logging.getLogger(__name__)

router = APIRouter()

# Platform caps teams at 8.
MAX_TEAM_SIZE = 8

PROJECT_SERVICE_URL = os.getenv("PROJECT_SERVICE_URL", "http://localhost:3002")


def _service_auth_secret() -> str:
    """Outgoing X-Service-Auth secret. Read at call time so tests can override."""
    return os.getenv("SERVICE_AUTH_SECRET", "")


# Stable keys the frontend maps to localized labels; do not rename lightly.
def _completeness_checks(messages: list) -> dict[str, bool]:
    """Map each BRD info requirement to whether the conversation covers it.

    Each key maps to a BRD template section that needs real data from the
    client. Empty conversation leaves every check False.
    """
    all_text = " ".join(m.content.lower() for m in messages if m.role == "user")

    return {
        # Section B — Executive Summary: project description present
        "description": len(all_text) > 80,
        # Section C — Problem Statement: pain points or motivation
        "problem": any(w in all_text for w in [
            "masalah", "problem", "kendala", "pain", "isu", "issue",
            "saat ini", "currently", "manual", "tidak bisa", "belum ada",
        ]),
        # Section D — Business Objectives: goals
        "objectives": any(w in all_text for w in [
            "tujuan", "goal", "objective", "target", "ingin", "mau", "want",
            "meningkatkan", "increase", "menurunkan", "reduce",
        ]),
        # Section E — Scope: features (in-scope)
        "features": any(w in all_text for w in [
            "fitur", "feature", "fungsi", "function", "modul", "module",
            "halaman", "page", "dashboard", "login", "register",
        ]),
        # Section G — Target Users
        "users": any(w in all_text for w in [
            "user", "pengguna", "pelanggan", "customer", "target", "audience",
            "admin", "konsumen", "pembeli", "buyer",
        ]),
        # Section H — Business Needs: non-trivial requirement detail
        "requirements": len(all_text) > 300 and any(w in all_text for w in [
            "harus", "must", "perlu", "need", "require", "wajib",
            "sistem", "system", "data", "laporan", "report",
        ]),
        # Section K — Risks / Assumptions
        "risks": any(w in all_text for w in [
            "risiko", "risk", "asumsi", "assumption", "keterbatasan", "constraint",
            "tantangan", "challenge", "hambatan",
        ]),
        # Section L — Success Metrics
        "metrics": any(w in all_text for w in [
            "metrik", "metric", "kpi", "ukur", "measure", "sukses", "success",
            "persentase", "percent", "angka", "number", "target",
        ]),
        # Section M — Constraints: budget
        "budget": any(w in all_text for w in [
            "budget", "biaya", "harga", "anggaran", "rp", "juta", "ribu",
            "million", "cost", "dana",
        ]),
        # Section M — Constraints: timeline
        "timeline": any(w in all_text for w in [
            "deadline", "waktu", "timeline", "kapan", "bulan", "minggu",
            "hari", "day", "week", "month", "selesai", "launch",
        ]),
        # Integrations (enriches H and E)
        "integrations": any(w in all_text for w in [
            "integrasi", "integration", "api", "payment", "pembayaran",
            "whatsapp", "google", "midtrans", "xendit", "notifikasi",
        ]),
    }


def calculate_completeness(messages: list) -> int:
    """Score chat conversation 0-100 against BRD template info requirements."""
    checks = _completeness_checks(messages)
    return min(100, int(sum(checks.values()) / len(checks) * 100))


def identify_missing(messages: list) -> list[str]:
    """Keys of BRD info requirements the conversation has not covered yet."""
    return [key for key, covered in _completeness_checks(messages).items() if not covered]


def _score_brd_against_template(brd: dict) -> BrdTemplateScore:
    """Score generated BRD dict against KerjaCUS! BRD template sections (A-N).

    Template sections mapped to BrdDocument fields:
      B  Executive Summary     → executive_summary (length + substance)
      D  Business Objectives   → business_objectives (count + specificity)
      E  Scope                 → scope + out_of_scope
      H  Business Needs/Reqs  → functional_requirements + non_functional_requirements
      K  Risks & Assumptions   → risk_assessment
      L  Success Metrics       → success_metrics
      M  Constraints           → estimated_price_min/max + estimated_timeline_days
    Sections F (Stakeholders), G (Target Users), I (Business Rules),
    J (Expected Benefits), N (Timeline detail) are not in BrdDocument schema —
    marked as gaps with score 0.
    """

    def _score_text(val: object, min_len: int = 100) -> tuple[int, str]:
        if not val:
            return 0, "empty"
        text = str(val)
        if len(text) >= min_len * 2:
            return 100, f"{len(text)} chars"
        if len(text) >= min_len:
            return 70, f"{len(text)} chars (adequate)"
        return 40, f"{len(text)} chars (too brief)"

    def _score_list(val: object, min_items: int = 3, ideal: int = 5) -> tuple[int, str]:
        if not val or not isinstance(val, list):
            return 0, "empty"
        n = len(val)
        if n >= ideal:
            return 100, f"{n} items"
        if n >= min_items:
            return 70, f"{n} items (adequate, aim for {ideal}+)"
        if n >= 1:
            return 40, f"{n} item(s) (too few, need {min_items}+)"
        return 0, "empty list"

    sections: list[BrdSectionScore] = []

    # B — Executive Summary
    s, r = _score_text(brd.get("executive_summary"), min_len=150)
    sections.append(BrdSectionScore(section="B", label="Executive Summary", score=s, reason=r))

    # D — Business Objectives
    s, r = _score_list(brd.get("business_objectives"), min_items=4, ideal=6)
    sections.append(BrdSectionScore(section="D", label="Business Objectives", score=s, reason=r))

    # E — Scope (in-scope)
    s, r = _score_text(brd.get("scope"), min_len=80)
    sections.append(BrdSectionScore(section="E", label="Scope (In-Scope)", score=s, reason=r))

    # E — Scope (out-of-scope)
    s, r = _score_list(brd.get("out_of_scope"), min_items=3, ideal=5)
    sections.append(BrdSectionScore(section="E", label="Scope (Out-of-Scope)", score=s, reason=r))

    # F — Stakeholders (not in schema — always gap)
    sections.append(BrdSectionScore(
        section="F", label="Stakeholders & Roles",
        score=0, reason="Not captured in current BRD schema",
    ))

    # G — Target Users (not in schema — always gap)
    sections.append(BrdSectionScore(
        section="G", label="Target User Segments",
        score=0, reason="Not captured in current BRD schema",
    ))

    # H — Functional Requirements
    s, r = _score_list(brd.get("functional_requirements"), min_items=4, ideal=7)
    sections.append(BrdSectionScore(section="H", label="Functional Requirements", score=s, reason=r))

    # H — Non-Functional Requirements
    s, r = _score_list(brd.get("non_functional_requirements"), min_items=4, ideal=7)
    sections.append(BrdSectionScore(section="H", label="Non-Functional Requirements", score=s, reason=r))

    # I — Business Rules (not in schema — gap)
    sections.append(BrdSectionScore(
        section="I", label="Business Rules",
        score=0, reason="Not captured in current BRD schema",
    ))

    # J — Expected Benefits (not in schema — gap)
    sections.append(BrdSectionScore(
        section="J", label="Expected Benefits",
        score=0, reason="Not captured in current BRD schema",
    ))

    # K — Risks & Assumptions
    s, r = _score_list(brd.get("risk_assessment"), min_items=3, ideal=5)
    sections.append(BrdSectionScore(section="K", label="Risks & Assumptions", score=s, reason=r))

    # L — Success Metrics
    s, r = _score_list(brd.get("success_metrics"), min_items=3, ideal=5)
    sections.append(BrdSectionScore(section="L", label="Success Metrics", score=s, reason=r))

    # M — Budget constraint
    price_min = brd.get("estimated_price_min", 0)
    price_max = brd.get("estimated_price_max", 0)
    if price_min > 0 and price_max > 0:
        bm_score, bm_reason = 100, f"Rp {price_min:,} – Rp {price_max:,}"
    elif price_min > 0 or price_max > 0:
        bm_score, bm_reason = 60, "Partial budget range"
    else:
        bm_score, bm_reason = 0, "No budget estimate"
    sections.append(BrdSectionScore(section="M", label="Budget Estimate", score=bm_score, reason=bm_reason))

    # M — Timeline & team size constraint
    tl = brd.get("estimated_timeline_days", 0)
    ts = brd.get("estimated_team_size", 0)
    if tl > 0 and ts > 0:
        tl_score, tl_reason = 100, f"{tl} days, {ts} person(s)"
    elif tl > 0:
        tl_score, tl_reason = 60, f"{tl} days (team size missing)"
    else:
        tl_score, tl_reason = 0, "No timeline estimate"
    sections.append(BrdSectionScore(section="M", label="Timeline & Team Size", score=tl_score, reason=tl_reason))

    # N — High-level timeline phases (not in schema — gap)
    sections.append(BrdSectionScore(
        section="N", label="High-Level Timeline Phases",
        score=0, reason="Not captured in current BRD schema",
    ))

    total = sum(s.score for s in sections)
    overall = round(total / len(sections)) if sections else 0
    return BrdTemplateScore(overall=overall, sections=sections)


@router.post(
    "/chat",
    response_model=ChatResponse,
    dependencies=[Depends(require_service_auth)],
    responses={502: {"description": "AI gateway unreachable"}},
)
async def chat_completion(request: ChatRequest):
    """AI chatbot for project scoping follow-up. Enriches context via RAG over past BRDs."""
    system_text, messages_payload = await _build_chat_messages_with_rag(request)
    missing = identify_missing(request.messages)
    if missing:
        system_text += (
            "\n\nInformation still missing for a complete BRD: "
            + ", ".join(missing)
            + ". Ask the client focused questions to fill these gaps before "
            "suggesting BRD generation."
        )

    try:
        async with track("chatbot", project_id=request.project_id) as rec:
            text = await generate_text(
                system_text,
                messages_payload,
                temperature=0.7,
                max_output_tokens=2048,
                on_usage=rec,
            )
    except LLMError as e:
        raise HTTPException(status_code=502, detail=f"AI service error: {e}") from e

    completeness = calculate_completeness(request.messages)

    return ChatResponse(
        message={"role": "assistant", "content": text},
        completeness_score=completeness,
        suggest_generate_brd=completeness >= 80,
        missing=missing,
    )


async def _build_chat_messages_with_rag(request: ChatRequest) -> tuple[str, list[dict]]:
    """Split system content from messages; append RAG context to system field.

    Gemini takes system text via system_instruction, not as a message role.
    Returns (system_text, user_assistant_messages).
    """
    rag_context_blocks: list[str] = []
    last_user_msg = next(
        (m.content for m in reversed(request.messages) if m.role == "user"),
        "",
    )
    if last_user_msg:
        try:
            from app.services.rag import hybrid_search

            chunks = await hybrid_search(
                query=last_user_msg,
                table="brd_documents",
                content_field="content",
                top_k=4,
            )
            rag_context_blocks = [c["content"] for c in chunks if c.get("content")]
        except Exception as e:
            logger.warning("RAG retrieval failed in /chat: %s", e)

    system_parts: list[str] = [m.content for m in request.messages if m.role == "system"]
    if rag_context_blocks:
        context_text = "\n\n---\n\n".join(rag_context_blocks)
        system_parts.append(
            "Context from similar past projects (use to ground your "
            "follow-up questions; do not reveal verbatim):\n"
            f"{context_text}"
        )

    user_assistant_messages = [
        m.model_dump() for m in request.messages if m.role in ("user", "assistant")
    ]
    system_text = "\n\n".join(p for p in system_parts if p).strip()
    return system_text, user_assistant_messages


def _sse(data: dict) -> bytes:
    """Encode dict as Server-Sent Event data line."""
    return f"data: {json.dumps(data, ensure_ascii=False)}\n\n".encode("utf-8")


def _split_system_from_messages(messages: list[dict]) -> tuple[str, list[dict]]:
    """Pull system content into a separate string for Gemini system_instruction.

    Gemini takes system text via system_instruction, not as a message role.
    """
    system_parts = [m["content"] for m in messages if m.get("role") == "system"]
    user_assistant = [m for m in messages if m.get("role") in ("user", "assistant")]
    return "\n\n".join(p for p in system_parts if p).strip(), user_assistant


async def _stream_chat_tokens(
    request: ChatRequest,
    system_text: str,
    messages_payload: list[dict],
) -> AsyncIterator[bytes]:
    """Stream Vertex deltas as SSE, finishing with completeness metadata."""
    full_text = ""
    usage: LlmUsage | None = None
    started = perf_counter()

    def sink(reported: LlmUsage) -> None:
        nonlocal usage
        usage = reported

    async def record(status: InteractionStatus) -> None:
        await record_interaction(
            "chatbot",
            usage=usage,
            latency_ms=round((perf_counter() - started) * 1000),
            status=status,
            project_id=request.project_id,
        )

    # Not `track`: awaiting in a generator's
    # finally breaks when the reader disconnects.
    try:
        async for delta in stream_text(
            system_text,
            messages_payload,
            temperature=0.7,
            max_output_tokens=2048,
            on_usage=sink,
        ):
            if delta:
                full_text += delta
                yield _sse({"type": "token", "delta": delta})
    except LLMError as e:
        await record("error")
        yield _sse({"type": "error", "message": f"AI gateway error: {e}"})
        return

    await record("success")

    completeness = calculate_completeness(request.messages)
    yield _sse(
        {
            "type": "done",
            "full_text": full_text,
            "completeness_score": completeness,
            "suggest_generate_brd": completeness >= 80,
            "missing": identify_missing(request.messages),
        }
    )


@router.post(
    "/chat/stream",
    dependencies=[Depends(require_service_auth)],
    responses={
        200: {
            "content": {"text/event-stream": {"schema": {}}},
            "description": "Server-Sent Events token stream",
        },
        502: {"description": "AI gateway unreachable"},
    },
)
async def chat_stream(request: ChatRequest):
    """Server-Sent Events stream for chatbot tokens; terminal event carries completeness."""
    system_text, messages_payload = await _build_chat_messages_with_rag(request)
    missing = identify_missing(request.messages)
    if missing:
        system_text += (
            "\n\nInformation still missing for a complete BRD: "
            + ", ".join(missing)
            + ". Ask the client focused questions to fill these gaps before "
            "suggesting BRD generation."
        )
    return StreamingResponse(
        _stream_chat_tokens(request, system_text, messages_payload),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache, no-transform",
            "X-Accel-Buffering": "no",
            "Connection": "keep-alive",
        },
    )


def _language_directive(language: str) -> str:
    """Instruct the model which language to write the document in."""
    if language == "en":
        return "Write the entire document in English."
    return (
        "Write the entire document in Indonesian (Bahasa Indonesia). "
        "Keep established technical terms in English where they are normally used that way "
        "(e.g. frontend, backend, database, API, deployment, framework)."
    )


BRD_SYSTEM_PROMPT = """You are a senior business analyst at KerjaCUS!, a managed marketplace platform for digital projects in Indonesia. Your job is to generate a comprehensive Business Requirement Document (BRD) from the project scoping conversation.

Analyze the conversation history carefully and produce a structured BRD in JSON format with these exact fields:

{
  "executive_summary": "A 2-3 paragraph summary of the project, its goals, target users, and key value proposition.",
  "business_objectives": ["List of 4-6 specific, measurable business objectives"],
  "success_metrics": ["List of 3-5 KPIs to measure project success"],
  "scope": "Detailed paragraph describing what is included in the project scope.",
  "out_of_scope": ["List of 3-5 items explicitly excluded from scope"],
  "functional_requirements": [
    {"title": "Feature Category Name", "content": "Detailed description of the feature and its sub-features"}
  ],
  "non_functional_requirements": ["List of 5-8 NFRs covering performance, security, scalability, accessibility"],
  "estimated_price_min": <integer in IDR>,
  "estimated_price_max": <integer in IDR>,
  "estimated_timeline_days": <integer>,
  "estimated_team_size": <integer>,
  "risk_assessment": ["List of 3-5 key risks with their mitigation strategies, each as a single string in format 'Risk: ... | Mitigation: ...'"]
}

Guidelines:
- The document language is stated in the user message; follow it exactly.
- Write like an experienced professional, not an AI. Be concrete, direct, and grounded in the actual use case. No hype words, no filler, no poetic or flowery phrasing, no redundancy; every sentence carries information. Do not invent facts that were not discussed.
- The executive summary is what persuades the project owner to proceed: make it compelling and business-focused, naming the value proposition, target users, and the return they can expect -- concrete and credible, never hype or filler.
- Be specific and actionable in requirements -- avoid vague statements.
- Price estimates should be realistic for the Indonesian market (developer rates Rp 15-40 million/month).
- Timeline should account for development, testing, and deployment.
- Team size should match the project complexity and timeline.
- Functional requirements should have 4-8 items covering all major feature areas.
- Always return valid JSON only, no markdown formatting or extra text."""


def _build_brd_messages(
    request: GenerateBrdRequest,
) -> list[dict]:
    """Build the message list for BRD generation from conversation history."""
    messages: list[dict] = [{"role": "system", "content": BRD_SYSTEM_PROMPT}]

    # Add conversation context as a consolidated user message
    conversation_text_parts: list[str] = []
    for msg in request.conversation_history:
        prefix = "Client" if msg.role == "user" else "AI Assistant"
        conversation_text_parts.append(f"{prefix}: {msg.content}")

    context_parts = [
        f"Project Category: {request.project_category}",
    ]
    if request.budget_min is not None:
        context_parts.append(f"Budget Min: Rp {request.budget_min:,}")
    if request.budget_max is not None:
        context_parts.append(f"Budget Max: Rp {request.budget_max:,}")
    if request.timeline_days is not None:
        context_parts.append(f"Requested Timeline: {request.timeline_days} days")

    user_prompt = (
        "Generate a BRD based on the following project scoping conversation and metadata.\n\n"
        f"--- Language ---\n{_language_directive(request.language)}\n\n"
        f"--- Project Metadata ---\n{chr(10).join(context_parts)}\n\n"
        f"--- Scoping Conversation ---\n{chr(10).join(conversation_text_parts)}\n\n"
    )
    if request.revision_instruction:
        user_prompt += (
            "--- Current BRD (revise this) ---\n"
            f"{json.dumps(request.current_document, indent=2, default=str)}\n\n"
            f"--- Revision Instruction ---\n{request.revision_instruction}\n\n"
            "Apply the revision instruction to the current BRD. Keep the rest intact and "
            "return the full updated BRD.\n\n"
        )
    user_prompt += "Return ONLY valid JSON matching the schema described in the system prompt."
    messages.append({"role": "user", "content": user_prompt})
    return messages


def _build_fallback_brd(request: GenerateBrdRequest) -> dict:
    """Build a reasonable BRD from request metadata when LLM fails."""
    # Extract project context from conversation
    conversation_text = " ".join(
        m.content for m in request.conversation_history if m.role == "user"
    )
    summary = conversation_text[:600] if conversation_text else "Digital project"

    budget_min = request.budget_min or 10_000_000
    budget_max = request.budget_max or 50_000_000
    timeline = request.timeline_days or 60
    team_size = max(1, min(MAX_TEAM_SIZE, timeline // 30))

    category_label = request.project_category.replace("_", " ").title()

    return {
        "executive_summary": (
            f"This {category_label} project aims to deliver a digital solution "
            f"based on the client's requirements gathered during scoping. "
            f"The project targets completion within {timeline} days with an estimated "
            f"budget range of Rp {budget_min:,} to Rp {budget_max:,}. "
            f"Key details: {summary[:300]}"
        ),
        "business_objectives": [
            f"Deliver a functional {category_label} within {timeline} days",
            "Meet all core functional requirements defined in scope",
            "Achieve responsive design supporting mobile and desktop",
            "Integrate with required third-party services",
        ],
        "success_metrics": [
            "All defined milestones completed on schedule",
            "Client approval on each milestone deliverable",
            "System passes functional and performance testing",
        ],
        "scope": (
            f"Full development of a {category_label} including frontend, backend, "
            "database design, API development, third-party integrations, "
            "testing, and deployment. Details based on scoping conversation."
        ),
        "out_of_scope": [
            "Native mobile applications (unless specified)",
            "Ongoing maintenance and support post-delivery",
            "Content creation and data migration",
        ],
        "functional_requirements": [
            {
                "title": "Core Application Features",
                "content": "Primary features as discussed during project scoping.",
            },
            {
                "title": "User Management",
                "content": "User registration, authentication, and profile management.",
            },
            {
                "title": "Admin Dashboard",
                "content": "Administrative interface for content and user management.",
            },
        ],
        "non_functional_requirements": [
            "Page load time under 3 seconds on 4G connections",
            "Support for 100+ concurrent users",
            "HTTPS encryption for all data in transit",
            "Responsive design for mobile and desktop",
            "Basic SEO optimization",
        ],
        "estimated_price_min": budget_min,
        "estimated_price_max": budget_max,
        "estimated_timeline_days": timeline,
        "estimated_team_size": team_size,
        "risk_assessment": [
            "Risk: Scope creep from additional requirements | Mitigation: Strict change request process with impact analysis",
            "Risk: Third-party API integration delays | Mitigation: Begin integration early, prepare fallback options",
            "Risk: Timeline pressure affecting quality | Mitigation: Prioritize core features, defer nice-to-haves to Phase 2",
        ],
        "language": request.language,
    }


def _parse_brd_response(parsed: dict, request: GenerateBrdRequest) -> dict:
    """Shape the LLM JSON into a BRD, fallback when empty."""
    if not parsed:
        logger.error(
            "BRD fell back to template: no JSON in the model response, project=%s",
            request.project_id,
        )
        return _build_fallback_brd(request)

    # Normalize functional_requirements: LLM may return {title, content} or {title, description}
    raw_reqs = parsed.get("functional_requirements", [])
    normalized_reqs = []
    for req in raw_reqs:
        if isinstance(req, dict):
            normalized_reqs.append(
                {
                    "title": req.get("title", "Feature"),
                    "content": req.get("content") or req.get("description", ""),
                }
            )
        elif isinstance(req, str):
            normalized_reqs.append({"title": "Requirement", "content": req})

    # Normalize risk_assessment: accept both string list and object list
    raw_risks = parsed.get("risk_assessment", [])
    normalized_risks = []
    for risk in raw_risks:
        if isinstance(risk, str):
            normalized_risks.append(risk)
        elif isinstance(risk, dict):
            r = risk.get("risk", risk.get("title", ""))
            m = risk.get("mitigation", risk.get("strategy", ""))
            normalized_risks.append(f"Risk: {r} | Mitigation: {m}" if m else r)

    fallback = _build_fallback_brd(request)

    return {
        "executive_summary": parsed.get("executive_summary") or fallback["executive_summary"],
        "business_objectives": parsed.get("business_objectives") or fallback["business_objectives"],
        "success_metrics": parsed.get("success_metrics") or fallback["success_metrics"],
        "scope": parsed.get("scope") or fallback["scope"],
        "out_of_scope": parsed.get("out_of_scope") or fallback["out_of_scope"],
        "functional_requirements": normalized_reqs or fallback["functional_requirements"],
        "non_functional_requirements": parsed.get("non_functional_requirements") or fallback["non_functional_requirements"],
        "estimated_price_min": parsed.get("estimated_price_min") or fallback["estimated_price_min"],
        "estimated_price_max": parsed.get("estimated_price_max") or fallback["estimated_price_max"],
        "estimated_timeline_days": parsed.get("estimated_timeline_days") or fallback["estimated_timeline_days"],
        "estimated_team_size": parsed.get("estimated_team_size") or fallback["estimated_team_size"],
        "risk_assessment": normalized_risks or fallback["risk_assessment"],
        # The owner picks the language, not the model.
        "language": request.language,
    }


@router.post(
    "/generate-brd",
    response_model=GenerateBrdResponse,
    dependencies=[Depends(require_service_auth)],
    responses={502: {"description": "AI gateway unreachable"}},
)
async def generate_brd(request: GenerateBrdRequest):
    """Generate BRD from conversation history via Vertex express."""
    messages = _build_brd_messages(request)
    model_used = "gemini-2.5-flash"
    tokens_used = 0

    async with track("brd_generation", project_id=request.project_id) as rec:
        try:
            system_text, user_msgs = _split_system_from_messages(messages)
            result = await generate_json(
                system_text,
                user_msgs,
                temperature=0.3,
                max_output_tokens=8192,
                on_usage=rec,
            )
            tokens_used = result.tokens
            model_used = result.model
            brd = _parse_brd_response(result.data, request)

        except LLMError as exc:
            rec.record_failure(exc)
            logger.error(
                "BRD fell back to template: gateway call failed, project=%s: %s",
                request.project_id,
                exc,
            )
            brd = _build_fallback_brd(request)

    template_score = _score_brd_against_template(brd)

    await publish_event(
        "ai.brd.generated",
        {
            "projectId": request.project_id,
            "tokensUsed": tokens_used,
            "model": model_used,
            "templateScore": template_score.overall,
        },
    )

    return GenerateBrdResponse(
        brd=brd,
        tokens_used=tokens_used,
        model=model_used,
        template_score=template_score,
    )


PRD_SYSTEM_PROMPT = """You are a senior technical architect at KerjaCUS!, a managed marketplace platform for digital projects in Indonesia. Your job is to generate a comprehensive Product Requirement Document (PRD) from the BRD and project context.

Analyze the BRD content and conversation history carefully and produce a structured PRD in JSON format with these exact fields:

{
  "tech_stack": ["List of recommended technologies, e.g. React, Node.js, PostgreSQL"],
  "architecture": "Detailed paragraph describing the system architecture (microservices, monolith, serverless, etc.) with justification.",
  "api_design": "Description of the API design approach (REST, GraphQL, etc.) with key endpoints listed.",
  "database_schema": "Description of the database design with key tables, relationships, and indexing strategy.",
  "team_composition": {
    "team_size": <integer>,
    "work_packages": [
      {
        "title": "Work Package Name (e.g. Frontend Development)",
        "description": "Detailed description of the work package scope",
        "required_skills": ["skill1", "skill2"],
        "estimated_hours": <float>,
        "amount": <integer in IDR>,
        "deliverables": [
          {"title": "Concrete output name, e.g. 'Checkout UI'", "type": "code | document | file | demo", "expected": "What a complete, acceptable version looks like"}
        ],
        "acceptance_criteria": ["Verifiable, testable statement the owner checks to accept the work, e.g. 'Checkout passes on mobile and desktop with saved addresses'"]
      }
    ]
  },
  "work_packages": [<same as team_composition.work_packages>],
  "sprint_plan": [
    {
      "sprint_number": <integer>,
      "title": "Sprint Title",
      "tasks": ["Task 1 description", "Task 2 description"],
      "duration_days": <integer, typically 14>
    }
  ],
  "dependencies": [
    {
      "from_package": "Work Package Title that must finish first",
      "to_package": "Work Package Title that depends on it",
      "type": "finish_to_start"
    }
  ],
  "assumptions": ["A condition the plan depends on being true, e.g. 'Owner supplies brand assets before sprint 1'"],
  "risks": ["Risk: ... | Mitigation: ..."],
  "estimated_price_min": <integer in IDR>,
  "estimated_price_max": <integer in IDR>,
  "estimated_timeline_days": <integer>,
  "estimated_team_size": <integer>
}

Guidelines:
- The document language is stated in the user message; follow it exactly. Code, identifiers, and established technical terms stay in English regardless.
- Write like an experienced engineer, not an AI. Be concrete, direct, and grounded in the actual use case. No hype words, no filler, no poetic phrasing, no redundancy; every sentence carries information. Do not invent requirements that were not in the BRD.
- This PRD is the brief an assigned talent builds from: every work package must be concrete enough to execute without further clarification.
- Each work package MUST list its deliverables (typed as code, document, file, or demo) and acceptance_criteria that are verifiable and testable -- the exact checks the owner runs to accept the work. No vague "works well".
- Tech stack should be specific (versions if relevant) and justified for the project type.
- Team size calculation: total_estimated_hours / (timeline_days * 6 working_hours_per_day), minimum 1, maximum 8.
- Work packages should be decomposed by role/skill area (Frontend, Backend, UI/UX, etc.).
- Sprint plan should have 2-week sprints covering the full timeline.
- Dependencies should form a valid DAG (no cycles).
- assumptions state what must hold for the plan to work; risks name concrete technical or delivery risks with a mitigation.
- Pricing should be realistic for the Indonesian market.
- Always return valid JSON only, no markdown formatting or extra text."""


def _build_prd_messages(request: GeneratePrdRequest) -> list[dict]:
    """Build the message list for PRD generation from BRD and conversation."""
    messages: list[dict] = [{"role": "system", "content": PRD_SYSTEM_PROMPT}]

    context_parts = [
        f"Project Category: {request.project_category}",
    ]
    if request.budget_min is not None:
        context_parts.append(f"Budget Min: Rp {request.budget_min:,}")
    if request.budget_max is not None:
        context_parts.append(f"Budget Max: Rp {request.budget_max:,}")
    if request.timeline_days is not None:
        context_parts.append(f"Requested Timeline: {request.timeline_days} days")

    conversation_text_parts: list[str] = []
    for msg in request.conversation_history:
        prefix = "Client" if msg.role == "user" else "AI Assistant"
        conversation_text_parts.append(f"{prefix}: {msg.content}")

    brd_json = json.dumps(request.brd_content, indent=2, default=str)

    user_prompt = (
        "Generate a PRD based on the following BRD document and project metadata.\n\n"
        f"--- Language ---\n{_language_directive(request.language)}\n\n"
        f"--- Project Metadata ---\n{chr(10).join(context_parts)}\n\n"
        f"--- BRD Document ---\n{brd_json}\n\n"
    )
    if conversation_text_parts:
        user_prompt += f"--- Scoping Conversation ---\n{chr(10).join(conversation_text_parts)}\n\n"
    if request.revision_instruction:
        user_prompt += (
            "--- Current PRD (revise this) ---\n"
            f"{json.dumps(request.current_document, indent=2, default=str)}\n\n"
            f"--- Revision Instruction ---\n{request.revision_instruction}\n\n"
            "Apply the revision instruction to the current PRD. Keep the rest intact and "
            "return the full updated PRD.\n\n"
        )
    user_prompt += "Return ONLY valid JSON matching the schema described in the system prompt."

    messages.append({"role": "user", "content": user_prompt})
    return messages


def _build_fallback_prd(request: GeneratePrdRequest) -> dict:
    """Build a reasonable PRD from BRD data when LLM fails."""
    brd = request.brd_content
    budget_min = request.budget_min or brd.get("estimated_price_min", 10_000_000)
    budget_max = request.budget_max or brd.get("estimated_price_max", 50_000_000)
    timeline = request.timeline_days or brd.get("estimated_timeline_days", 60)
    team_size = brd.get("estimated_team_size", max(1, min(MAX_TEAM_SIZE, timeline // 30)))

    category_label = request.project_category.replace("_", " ").title()

    # Default work packages based on category
    work_packages = [
        {
            "title": "Backend API Development",
            "description": f"Server-side logic, API endpoints, database integration for {category_label}",
            "required_skills": ["Node.js", "PostgreSQL", "REST API"],
            "estimated_hours": float(timeline * 4),
            "amount": int(budget_min * 0.35),
            "deliverables": [
                {
                    "title": "REST API endpoints",
                    "type": "code",
                    "expected": "Every endpoint in the PRD implemented with input validation",
                },
                {
                    "title": "API documentation",
                    "type": "document",
                    "expected": "OpenAPI 3.1 spec covering every endpoint",
                },
            ],
            "acceptance_criteria": [
                "All endpoints return the documented responses and reject invalid input",
                "Integration tests cover the main flows and pass",
            ],
        },
        {
            "title": "Frontend Development",
            "description": f"User interface implementation for {category_label}",
            "required_skills": ["React", "TypeScript", "Tailwind CSS"],
            "estimated_hours": float(timeline * 4),
            "amount": int(budget_min * 0.35),
            "deliverables": [
                {
                    "title": "UI implementation",
                    "type": "code",
                    "expected": "All screens built to the approved design and wired to the API",
                },
            ],
            "acceptance_criteria": [
                "The UI matches the approved design on mobile and desktop",
                "Forms validate and submit correctly against the API",
            ],
        },
        {
            "title": "UI/UX Design",
            "description": "Wireframes, mockups, design system, and prototypes",
            "required_skills": ["Figma", "UI Design", "UX Research"],
            "estimated_hours": float(timeline * 2),
            "amount": int(budget_min * 0.2),
            "deliverables": [
                {
                    "title": "Figma design file",
                    "type": "file",
                    "expected": "Wireframes, high-fidelity mockups, and a reusable design system",
                },
            ],
            "acceptance_criteria": [
                "The owner approves the high-fidelity mockups",
                "The design system covers every component used in the screens",
            ],
        },
    ]

    sprints = []
    num_sprints = max(1, timeline // 14)
    for i in range(num_sprints):
        sprints.append(
            {
                "sprint_number": i + 1,
                "title": f"Sprint {i + 1}",
                "tasks": [f"Development tasks for sprint {i + 1}"],
                "duration_days": 14,
            }
        )

    return {
        "tech_stack": ["React", "TypeScript", "Node.js", "PostgreSQL", "Tailwind CSS", "Docker"],
        "architecture": (
            f"Modular monolith architecture for {category_label} with clear service boundaries. "
            "REST API backend with PostgreSQL database, React frontend with server-side rendering support."
        ),
        "api_design": (
            "RESTful API design with versioned endpoints (/api/v1/*). "
            "JSON request/response format, JWT authentication, pagination for list endpoints."
        ),
        "database_schema": (
            "Normalized PostgreSQL schema with UUID primary keys, timestamptz for all timestamps. "
            "Key tables based on BRD functional requirements with proper indexing and foreign key constraints."
        ),
        "team_composition": {
            "team_size": team_size,
            "work_packages": work_packages,
        },
        "work_packages": work_packages,
        "sprint_plan": sprints,
        "dependencies": [
            {
                "from_package": "UI/UX Design",
                "to_package": "Frontend Development",
                "type": "finish_to_start",
            },
            {
                "from_package": "Backend API Development",
                "to_package": "Frontend Development",
                "type": "start_to_start",
            },
        ],
        "assumptions": [
            "The owner supplies branding, content, and any third-party credentials before the sprint that needs them",
            "Requirements are frozen at PRD approval; later changes go through the revision process",
        ],
        "risks": [
            "Risk: Scope creep from new requirements | Mitigation: Change requests are re-estimated before work starts",
            "Risk: Third-party integration delays | Mitigation: Integrate early and keep a fallback path",
        ],
        "estimated_price_min": budget_min,
        "estimated_price_max": budget_max,
        "estimated_timeline_days": timeline,
        "estimated_team_size": team_size,
        "language": request.language,
    }


def _norm_str_list(raw: object) -> list[str]:
    """Keep only the string entries of a list."""
    return [x for x in raw if isinstance(x, str)] if isinstance(raw, list) else []


def _norm_deliverables(raw: object) -> list[dict]:
    """Shape deliverables to {title, type, expected}, tolerating bare strings."""
    out: list[dict] = []
    for d in raw if isinstance(raw, list) else []:
        if isinstance(d, dict):
            out.append(
                {
                    "title": str(d.get("title", "")),
                    "type": str(d.get("type", "document")),
                    "expected": str(d.get("expected", "")),
                }
            )
        elif isinstance(d, str):
            out.append({"title": d, "type": "document", "expected": ""})
    return out


def _parse_prd_response(parsed: dict, request: GeneratePrdRequest) -> dict:
    """Shape the LLM JSON into a PRD, fallback when empty."""
    if not parsed:
        logger.error(
            "PRD fell back to template: no JSON in the model response, project=%s",
            request.project_id,
        )
        return _build_fallback_prd(request)

    fallback = _build_fallback_prd(request)

    # Normalize work_packages
    raw_wps = parsed.get("work_packages", [])
    normalized_wps = []
    for wp in raw_wps:
        if isinstance(wp, dict):
            normalized_wps.append(
                {
                    "title": wp.get("title", "Work Package"),
                    "description": wp.get("description", ""),
                    "required_skills": wp.get("required_skills", []),
                    "estimated_hours": float(wp.get("estimated_hours", 0)),
                    "amount": int(wp.get("amount", 0)),
                    "deliverables": _norm_deliverables(wp.get("deliverables")),
                    "acceptance_criteria": _norm_str_list(wp.get("acceptance_criteria")),
                }
            )

    # A work package with no amount or hours is dropped downstream: the project
    # service cannot create an unpriced package (amount/hours CHECK > 0), so
    # matching would find nothing to assign and confirm would dead-end. The LLM
    # sometimes names packages without pricing them, so backfill the gaps from
    # the same budget and timeline the fallback uses, keeping its decomposition.
    price_ref = int(fallback["estimated_price_min"])
    hours_ref = float(fallback["estimated_timeline_days"] * 4)
    share = max(1, int(price_ref * 0.9 / (len(normalized_wps) or 1)))
    for wp in normalized_wps:
        if wp["amount"] <= 0:
            wp["amount"] = share
        if wp["estimated_hours"] <= 0:
            wp["estimated_hours"] = hours_ref

    # Normalize sprint_plan
    raw_sprints = parsed.get("sprint_plan", [])
    normalized_sprints = []
    for sp in raw_sprints:
        if isinstance(sp, dict):
            normalized_sprints.append(
                {
                    "sprint_number": int(sp.get("sprint_number", len(normalized_sprints) + 1)),
                    "title": sp.get("title", f"Sprint {len(normalized_sprints) + 1}"),
                    "tasks": sp.get("tasks", []),
                    "duration_days": int(sp.get("duration_days", 14)),
                }
            )

    # Normalize dependencies
    raw_deps = parsed.get("dependencies", [])
    normalized_deps = []
    for dep in raw_deps:
        if isinstance(dep, dict):
            normalized_deps.append(
                {
                    "from_package": dep.get("from_package", ""),
                    "to_package": dep.get("to_package", ""),
                    "type": dep.get("type", "finish_to_start"),
                }
            )

    # Normalize team_composition
    raw_tc = parsed.get("team_composition", {})
    team_composition = {
        "team_size": raw_tc.get("team_size", parsed.get("estimated_team_size", fallback["estimated_team_size"])),
        "work_packages": normalized_wps or fallback["work_packages"],
    }

    return {
        "tech_stack": parsed.get("tech_stack") or fallback["tech_stack"],
        "architecture": parsed.get("architecture") or fallback["architecture"],
        "api_design": parsed.get("api_design") or fallback["api_design"],
        "database_schema": parsed.get("database_schema") or fallback["database_schema"],
        "team_composition": team_composition,
        "work_packages": normalized_wps or fallback["work_packages"],
        "sprint_plan": normalized_sprints or fallback["sprint_plan"],
        "dependencies": normalized_deps or fallback["dependencies"],
        "assumptions": _norm_str_list(parsed.get("assumptions")) or fallback["assumptions"],
        "risks": _norm_str_list(parsed.get("risks")) or fallback["risks"],
        "estimated_price_min": parsed.get("estimated_price_min") or fallback["estimated_price_min"],
        "estimated_price_max": parsed.get("estimated_price_max") or fallback["estimated_price_max"],
        "estimated_timeline_days": parsed.get("estimated_timeline_days") or fallback["estimated_timeline_days"],
        "estimated_team_size": parsed.get("estimated_team_size") or fallback["estimated_team_size"],
        # The owner picks the language, not the model.
        "language": request.language,
    }


@router.post(
    "/generate-prd",
    response_model=GeneratePrdResponse,
    dependencies=[Depends(require_service_auth)],
    responses={502: {"description": "AI gateway unreachable"}},
)
async def generate_prd(request: GeneratePrdRequest):
    """Generate PRD from BRD content and conversation history via Vertex express."""
    messages = _build_prd_messages(request)
    model_used = "gemini-2.5-flash"
    tokens_used = 0

    async with track("prd_generation", project_id=request.project_id) as rec:
        try:
            system_text, user_msgs = _split_system_from_messages(messages)
            result = await generate_json(
                system_text,
                user_msgs,
                temperature=0.3,
                max_output_tokens=16384,
                on_usage=rec,
            )
            tokens_used = result.tokens
            model_used = result.model
            prd = _parse_prd_response(result.data, request)

        except LLMError as exc:
            rec.record_failure(exc)
            logger.error(
                "PRD fell back to template: gateway call failed, project=%s: %s",
                request.project_id,
                exc,
            )
            prd = _build_fallback_prd(request)

    await publish_event(
        "ai.prd.generated",
        {
            "projectId": request.project_id,
            "tokensUsed": tokens_used,
            "model": model_used,
        },
    )

    return GeneratePrdResponse(
        prd=prd,
        tokens_used=tokens_used,
        model=model_used,
    )


def _resolve_cv_source_url(raw_file_url: str) -> str:
    """Resolve a CV reference to a URL on our own storage, refusing anything else.

    Every real call passes an S3 object key: the browser uploads via a presigned
    PUT and posts back `presigned.key`. This handler used to fetch any absolute
    http(s) URL verbatim and return 2000 bytes of the response body to the
    caller, and it is the one route without service auth - so it was an
    unauthenticated full-read SSRF reaching cloud metadata, internal services
    and anything else the container could route to.

    Absolute URLs are now accepted only for our configured storage hosts, and
    keys may not traverse out of the bucket.
    """
    # 400, not 422: FastAPI reserves 422 for HTTPValidationError, whose `detail`
    # is an array. Returning a string there violates the published schema.
    key = (raw_file_url or "").strip()
    if not key:
        raise HTTPException(status_code=400, detail="file_url is required")

    s3_endpoint = os.getenv("S3_ENDPOINT", "http://localhost:9000")
    bucket = os.getenv("S3_BUCKET", "kerjacus-uploads")
    public_url = os.getenv("S3_PUBLIC_URL", "")

    if key.startswith(("http://", "https://")):
        allowed = {urlparse(u).netloc for u in (s3_endpoint, public_url) if u}
        if urlparse(key).netloc not in allowed:
            raise HTTPException(status_code=403, detail="file_url must reference project storage")
        return key

    if ".." in key:
        raise HTTPException(status_code=403, detail="file_url must reference project storage")

    return f"{s3_endpoint.rstrip('/')}/{bucket}/{key.lstrip('/')}"


@router.post(
    "/parse-cv",
    response_model=CvParseResponse,
    # Returns name, email, phone and history. Callers must be services.
    dependencies=[Depends(require_service_auth)],
    responses={
        401: {"description": "missing or wrong X-Service-Auth"},
        403: {"description": "file_url does not reference project storage"},
        502: {"description": "CV could not be downloaded for parsing"},
    },
)
async def parse_cv(request: CvParseRequest):
    """Parse CV using document text extraction + Vertex structured extraction."""
    import asyncio
    import tempfile
    from pathlib import Path

    from pydantic import BaseModel, Field

    class ExtractedCV(BaseModel):
        name: str = Field(default="", description="Full name")
        email: str = Field(default="", description="Email address")
        phone: str = Field(default="", description="Phone number")
        summary: str = Field(default="", description="Professional summary or objective statement")
        skills: list[str] = Field(
            default_factory=list,
            description=(
                "ALL technical skills extracted from EVERY section: "
                "certifications tech tags, project tech stacks, work experience descriptions, "
                "education coursework, and any explicit skills section. "
                "Include frameworks, languages, tools, platforms, algorithms, and ML model types."
            ),
        )
        # Typed entries, not list[dict]: Instructor turns this model into the
        # JSON schema the LLM is constrained by, so nested fields with their own
        # descriptions are what actually drive extraction detail - and give the
        # retry loop something to validate against.
        education: list[EducationEntry] = Field(
            default_factory=list,
            description="Every education entry, most recent first",
        )
        experience: list[ExperienceEntry] = Field(
            default_factory=list,
            description="Paid professional roles, internships included",
        )
        organizational_experience: list[ExperienceEntry] = Field(
            default_factory=list,
            description=(
                "Unpaid roles: student organisations, committees, volunteering. "
                "Use the organisation as company and the role held as position"
            ),
        )
        projects: list[ProjectEntry] = Field(
            default_factory=list,
            description="Personal or academic projects",
        )
        certifications: list[CertificationEntry] = Field(
            default_factory=list,
            description="Certifications, courses and professional training",
        )
        portfolio_urls: list[str] = Field(
            default_factory=list,
            description="Portfolio/professional URLs (GitHub, LinkedIn, Dribbble, Behance, etc.)",
        )
        years_of_experience: int | None = Field(
            default=None,
            description="Total years of professional work experience (integer, exclude internships if under 1 year)",
        )

    file_url = _resolve_cv_source_url(request.file_url)

    file_bytes = None
    for attempt in range(3):
        try:
            async with httpx.AsyncClient(timeout=15.0) as dl:
                res = await dl.get(file_url)
                if res.status_code == 200:
                    file_bytes = res.content
                    break
                logger.warning(
                    "CV download attempt %d failed: status=%d url=%s",
                    attempt + 1, res.status_code, file_url,
                )
        except Exception as e:
            logger.warning("CV download attempt %d errored: %s", attempt + 1, e)
        await asyncio.sleep(1)

    # A transport failure is not a bad CV; surface it so the caller can retry
    # rather than persisting a fake zero-confidence, unverified result.
    if file_bytes is None:
        raise HTTPException(
            status_code=502,
            detail="Could not download the CV for parsing; please try again",
        )

    # Step 2: Extract text based on file type
    cv_text = ""

    if file_bytes:
        ext = (request.file_type or "pdf").lower()
        try:
            if ext == "pdf":
                try:
                    import pypdfium2 as pdfium
                    pdf = pdfium.PdfDocument(file_bytes)
                    pages = []
                    for page in pdf:
                        textpage = page.get_textpage()
                        pages.append(textpage.get_text_bounded())
                        textpage.close()
                        page.close()
                    pdf.close()
                    cv_text = "\n".join(pages)
                except Exception:
                    cv_text = file_bytes.decode("utf-8", errors="ignore")
            elif ext in ("docx", "doc"):
                tmp_path = None
                try:
                    import docx
                    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                        tmp.write(file_bytes)
                        tmp_path = tmp.name
                    doc = docx.Document(tmp_path)
                    cv_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                except Exception:
                    cv_text = file_bytes.decode("utf-8", errors="ignore")
                finally:
                    if tmp_path:
                        Path(tmp_path).unlink(missing_ok=True)
            else:
                cv_text = file_bytes.decode("utf-8", errors="ignore")
        except Exception:
            cv_text = file_bytes.decode("utf-8", errors="ignore")

    if not cv_text or len(cv_text.strip()) < 50:
        return CvParseResponse(
            talent_id=request.talent_id,
            parsed_data=CvParsedData(),
            confidence_score=0.0,
            raw_text="",
        )

    # Step 3: Vertex structured extraction
    cv_system_prompt = (
        "You are an expert CV parser. Extract ALL structured information from this CV/resume text.\n\n"
        "CRITICAL — skills extraction rules:\n"
        "1. Scan EVERY section: certifications tech tags, project tech stacks (after '|' or in parentheses), "
        "work experience bullet points, education coursework, and any explicit skills section.\n"
        "2. Include: programming languages, frameworks, libraries, tools, platforms, cloud services, "
        "databases, ML algorithms (XGBoost, CatBoost, LightGBM, KNN, GNB, CNN, LSTM, etc.), "
        "MLOps tools (MLflow, Kubeflow, KServe, Feast), data tools (Tableau, R, Streamlit, Gradio), "
        "AI frameworks (LangChain, FAISS, Transformers, Hugging Face, LLM).\n"
        "3. Do NOT invent skills not mentioned in the text.\n"
        "4. Deduplicate — return each skill once.\n\n"
        "For organizational_experience: extract leadership roles in student orgs, volunteer work, committees.\n"
        "For projects: extract from 'Projects' section — each with title, tech stack list, short description, and any URL.\n"
        "For certifications: include the tech tags listed after the cert name (e.g. 'IBM AI Engineering | Python, PyTorch, TensorFlow').\n"
        "For Indonesian CVs, handle both Indonesian and English content."
    )

    async with track("cv_parsing") as rec:
        try:
            extracted = await generate_structured(
                cv_system_prompt,
                [{"role": "user", "content": cv_text[:12000]}],
                schema=ExtractedCV,
                temperature=0.1,
                max_output_tokens=4096,
                on_usage=rec,
            )

            parsed_data = CvParsedData(
                name=extracted.name,
                email=extracted.email,
                phone=extracted.phone,
                summary=extracted.summary or None,
                skills=extracted.skills,
                education=extracted.education,
                experience=extracted.experience,
                organizational_experience=extracted.organizational_experience,
                projects=extracted.projects,
                certifications=extracted.certifications,
                portfolio_urls=extracted.portfolio_urls,
                years_of_experience=extracted.years_of_experience,
            )

            # Confidence based on field completeness
            filled_fields = sum(1 for v in [
                extracted.name, extracted.email, extracted.phone,
                extracted.skills, extracted.education, extracted.experience,
            ] if v)
            confidence = min(0.95, 0.3 + (filled_fields / 6) * 0.65)

        except Exception as exc:
            # Retired model returns 404 here, silently before.
            rec.record_failure(exc)
            logger.error(
                "CV parsed by regex: LLM extraction failed, talent=%s: %s",
                request.talent_id,
                exc,
            )
            from app.services.cv_parser import parse_cv_text

            result = parse_cv_text(cv_text)
            parsed_data = CvParsedData(
                name=result.name,
                email=result.email,
                phone=result.phone,
                skills=result.skills,
                education=result.education,
                experience=result.experience,
                projects=result.projects,
                portfolio_urls=result.portfolio_urls,
            )
            # Score field completeness, not skill count alone. Counting only skills
            # meant a parse that recovered name, email, education and experience
            # scored no better than one that recovered a long list of noise - and it
            # rewarded the substring false-positives the skill matcher used to emit.
            # Capped below the Instructor path, which stays the more trusted source.
            fallback_fields = sum(
                1 for v in [
                    result.name, result.email, result.phone,
                    result.skills, result.education, result.experience,
                ] if v
            )
            confidence = min(0.7, 0.25 + (fallback_fields / 6) * 0.45)

    await publish_event(
        "ai.cv.parsed",
        {
            "talentId": request.talent_id,
            "confidenceScore": float(confidence),
            "skillCount": len(parsed_data.skills or []),
        },
    )

    return CvParseResponse(
        talent_id=request.talent_id,
        parsed_data=parsed_data,
        confidence_score=confidence,
        raw_text=cv_text[:2000],
    )


SPEC_PARSE_SYSTEM_PROMPT = """You are a project specification analyzer for KerjaCUS!, a managed marketplace for digital projects in Indonesia.
Extract key project information from the uploaded specification document.

Return a JSON object with exactly these fields:
{
  "summary": "A 2-3 sentence summary of the project",
  "features": ["list of key features or requirements mentioned"],
  "target_users": "description of intended users or audience",
  "integrations": ["list of third-party systems or APIs mentioned"],
  "tech_requirements": "any technical requirements, constraints, or preferences mentioned",
  "budget_hints": "any budget, cost, or pricing information mentioned (empty string if none)",
  "timeline_hints": "any timeline, deadline, or schedule information mentioned (empty string if none)",
  "completeness": <integer 0-100, how complete this spec is for generating a Business Requirements Document>
}

The completeness score should reflect how much information is available for generating a BRD:
- 90-100: Very detailed spec with features, users, tech, budget, timeline
- 70-89: Good spec with most key areas covered
- 50-69: Partial spec, missing several important areas
- 30-49: Brief overview, needs significant follow-up
- 0-29: Very sparse, barely any useful information

Return ONLY valid JSON, no markdown or extra text."""


@router.post(
    "/parse-spec",
    response_model=ParseSpecResponse,
    dependencies=[Depends(require_service_auth)],
)
async def parse_spec(request: ParseSpecRequest):
    """Parse an uploaded specification document and extract project information."""
    # Same storage allow-list as parse-cv: the owner controls file_url, and an
    # unrestricted fetch here was a server-side request forgery that reflected
    # the response body back to the caller.
    file_url = _resolve_cv_source_url(request.file_url)
    file_type = request.file_type
    notes = request.notes

    file_bytes = None
    try:
        async with httpx.AsyncClient(timeout=30.0) as client:
            res = await client.get(file_url)
            if res.status_code == 200:
                file_bytes = res.content
            else:
                logger.warning("Spec download failed: status=%d", res.status_code)
    except Exception as e:
        logger.warning("Spec download errored: %s", e)

    if not file_bytes:
        return ParseSpecResponse(
            data=ParseSpecData(
                summary="Failed to download specification file.",
                completeness=0,
            ),
        )

    # Parse document text
    from app.services.cv_parser import extract_text

    raw_text = extract_text(file_bytes, file_type)

    if not raw_text or len(raw_text.strip()) < 50:
        return ParseSpecResponse(
            data=ParseSpecData(
                summary="Document too short to extract meaningful information.",
                completeness=10,
            ),
        )

    # Extract project information via Vertex express
    async with track("spec_parsing") as rec:
        try:
            result = await generate_json(
                SPEC_PARSE_SYSTEM_PROMPT,
                [
                    {
                        "role": "user",
                        "content": f"Parse this specification document:\n\n{raw_text[:8000]}\n\nAdditional notes from the client: {notes}",
                    },
                ],
                temperature=0.2,
                max_output_tokens=1024,
                on_usage=rec,
            )
            parsed = result.data
            if parsed:
                return ParseSpecResponse(
                    data=ParseSpecData(
                        summary=parsed.get("summary", raw_text[:500]),
                        features=parsed.get("features", []),
                        target_users=parsed.get("target_users", ""),
                        integrations=parsed.get("integrations", []),
                        tech_requirements=parsed.get("tech_requirements", ""),
                        budget_hints=parsed.get("budget_hints", ""),
                        timeline_hints=parsed.get("timeline_hints", ""),
                        completeness=min(100, max(0, int(parsed.get("completeness", 50)))),
                    ),
                )
        except Exception as e:
            rec.record_failure(e)
            logger.warning("Spec LLM extraction failed, falling back to raw text: %s", e)

    # Fallback: return raw text summary
    return ParseSpecResponse(
        data=ParseSpecData(
            summary=raw_text[:500],
            completeness=40,
        ),
    )


@router.post(
    "/match-talents",
    response_model=MatchingResponse,
    dependencies=[Depends(require_service_auth)],
)
async def match_talents(request: MatchingRequest):
    """Match talents to a project: delegate scoring to project-service rule-based recommender."""
    headers = {}
    secret = _service_auth_secret()
    if secret:
        headers["X-Service-Auth"] = secret

    project_recommendations: list[dict] = []
    exploration_count = 0
    exploitation_count = 0
    try:
        async with httpx.AsyncClient(timeout=10.0) as client:
            res = await client.post(
                f"{PROJECT_SERVICE_URL}/api/v1/matching/recommend",
                json={
                    "requiredSkills": request.required_skills,
                    "limit": 10,
                },
                headers=headers,
            )
            res.raise_for_status()
            payload = res.json()
            data = payload.get("data", {}) if isinstance(payload, dict) else {}
            project_recommendations = data.get("recommendations", [])
            exploration_count = data.get("explorationCount", 0)
            exploitation_count = data.get("exploitationCount", 0)
    except httpx.HTTPError as e:
        logger.warning("project-service matching unavailable: %s", e)
        return MatchingResponse(
            project_id=request.project_id,
            recommendations=[],
            exploration_count=0,
            exploitation_count=0,
        )

    recommendations = []
    for r in project_recommendations:
        recommendations.append(
            {
                "talent_id": r.get("talentId", ""),
                "score": float(r.get("score", 0)),
                "skill_match": float(r.get("skillMatch", 0)),
                "pemerataan_score": float(r.get("pemerataanScore", 0)),
                "track_record": float(r.get("trackRecord", 0)),
                "rating": float(r.get("rating", 0)),
                "is_exploration": bool(r.get("isExploration", False)),
            }
        )

    await publish_event(
        "ai.matching.completed",
        {
            "projectId": request.project_id,
            "recommendationCount": len(recommendations),
            "explorationCount": exploration_count,
            "exploitationCount": exploitation_count,
        },
    )

    return MatchingResponse(
        project_id=request.project_id,
        recommendations=recommendations,
        exploration_count=exploration_count,
        exploitation_count=exploitation_count,
    )


class EmbedDocumentRequest(BaseModel):
    documentId: str
    documentType: Literal["brd", "prd"]
    content: str | dict | list


@router.post(
    "/embed-document",
    dependencies=[Depends(require_service_auth)],
    responses={
        500: {"description": "Embedding or persistence error"},
        503: {"description": "Embedding service unavailable"},
    },
)
async def embed_document(request: EmbedDocumentRequest):
    """Compute Gemini embedding for a BRD/PRD and persist it to the document row.

    Internal endpoint -- expected to be called from project-service after approval.
    """
    document_id = request.documentId
    document_type = request.documentType
    content = request.content

    if isinstance(content, dict | list):
        text_input = json.dumps(content, default=str)[:8000]
    else:
        text_input = str(content)[:8000]

    table = "brd_documents" if document_type == "brd" else "prd_documents"

    try:
        from app.services.embedding import embed_text
        from app.services.rag import write_embedding

        embedding = await embed_text(text_input)
        ok = await write_embedding(table=table, row_id=document_id, embedding=embedding)
        if not ok:
            raise HTTPException(status_code=500, detail="Failed to persist embedding")
    except HTTPException:
        raise
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=str(e)) from e
    except Exception as e:
        logger.exception("embed-document failed: %s", e)
        raise HTTPException(status_code=500, detail=f"Embedding error: {e}") from e

    return {"success": True, "documentId": document_id, "dimensions": len(embedding)}
